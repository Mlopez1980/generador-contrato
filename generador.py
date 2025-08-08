from docx import Document
from datetime import date
from num2words import num2words

PLANTILLA = "plantilla_promesa_marcadores.docx"

MODELOS = {
    "1E": {"area": 82.79, "habitaciones": "2", "ubicacion": "primer nivel de apartamentos, segundo nivel del edificio"},
    "2D": {"area": 82.79, "habitaciones": "2", "ubicacion": "segundo nivel de apartamentos, tercer nivel del edificio"},
    "2E": {"area": 82.79, "habitaciones": "2", "ubicacion": "segundo nivel de apartamentos, tercer nivel del edificio"},
    "2F": {"area": 44.55, "habitaciones": "1", "ubicacion": "segundo nivel de apartamentos, tercer nivel del edificio"},
    "3C": {"area": 45.997, "habitaciones": "1", "ubicacion": "tercer nivel de apartamentos, cuarto nivel del edificio"},
    "3E": {"area": 82.79, "habitaciones": "2", "ubicacion": "tercer nivel de apartamentos, cuarto nivel del edificio"},
    "3F": {"area": 44.55, "habitaciones": "1", "ubicacion": "tercer nivel de apartamentos, cuarto nivel del edificio"},
    "4D": {"area": 82.79, "habitaciones": "2", "ubicacion": "cuarto nivel de apartamentos, quinto nivel del edificio"},
    "4E": {"area": 82.79, "habitaciones": "2", "ubicacion": "cuarto nivel de apartamentos, quinto nivel del edificio"},
    "4F": {"area": 44.55, "habitaciones": "1", "ubicacion": "cuarto nivel de apartamentos, quinto nivel del edificio"},
    "5C": {"area": 45.997, "habitaciones": "1", "ubicacion": "quinto nivel de apartamentos, sexto nivel del edificio"},
    "5E": {"area": 82.79, "habitaciones": "2", "ubicacion": "quinto nivel de apartamentos, sexto nivel del edificio"},
    "5F": {"area": 44.55, "habitaciones": "1", "ubicacion": "quinto nivel de apartamentos, sexto nivel del edificio"}
}

DESCRIPCIONES = {
    "1": "dormitorio principal, baño completo",
    "2": "dormitorio principal con walk in closet y baño completo, dormitorio secundario, baño completo compartido"
}

def area_a_texto(area):
    parte_entera = int(area)
    parte_decimal = int(round((area - parte_entera) * 1000))
    texto_entera = num2words(parte_entera, lang='es')
    texto_decimal = num2words(parte_decimal, lang='es')
    return f"{texto_entera} punto {texto_decimal}"

def monto_a_texto_dolares(monto):
    try:
        monto = float(monto.replace(",", ""))
        parte_entera = int(monto)
        parte_decimal = int(round((monto - parte_entera) * 100))
        texto = num2words(parte_entera, lang='es')
        return f"{texto} dólares americanos con {parte_decimal:02d}/100"
    except:
        return "—"

def monto_formato_numerico(monto):
    try:
        monto = float(monto.replace(",", ""))
        return "{:,.2f}".format(monto)
    except:
        return monto

def reemplazar_marcadores(documento, datos):
    for parrafo in documento.paragraphs:
        for clave, valor in datos.items():
            if f"{{{{{clave}}}}}" in parrafo.text:
                parrafo.text = parrafo.text.replace(f"{{{{{clave}}}}}", valor)
    for tabla in documento.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                for clave, valor in datos.items():
                    if f"{{{{{clave}}}}}" in celda.text:
                        celda.text = celda.text.replace(f"{{{{{clave}}}}}", valor)

def generar_contrato_desde_formulario(entrada):
    modelo = entrada["MODELO"].upper()
    if modelo not in MODELOS:
        raise ValueError("Modelo inválido.")

    area = MODELOS[modelo]["area"]
    area_completa = f"{area_a_texto(area)} metros cuadrados ({area:.3f} M2)"
    habitaciones = MODELOS[modelo]["habitaciones"]
    descripcion = DESCRIPCIONES[habitaciones]

    precio = monto_formato_numerico(entrada["PRECIO"])
    reserva = monto_formato_numerico(entrada["MONTO_RESERVA"])
    complemento = monto_formato_numerico(entrada["COMPLEMENTO_PRIMA"])
    tercer_pago = monto_formato_numerico(entrada["TERCER_PAGO"])

    datos = {
        "NOMBRE": entrada["NOMBRE"],
        "DNI": entrada["DNI"],
        "ESTADO_CIVIL": entrada["ESTADO_CIVIL"],
        "PROFESION": entrada["PROFESION"],
        "FECHA": entrada["FECHA"] or date.today().strftime("%d/%m/%Y"),
        "PRECIO": precio,
        "PRECIO_TEXTO": monto_a_texto_dolares(entrada["PRECIO"]),
        "MODELO": modelo,
        "AREA": f"{area:.3f}",
        "AREA_COMPLETA": area_completa,
        "HABITACIONES": habitaciones,
        "UBICACION": MODELOS[modelo]["ubicacion"],
        "DESCRIPCION_HABITACIONES": descripcion,
        "MONTO_RESERVA": reserva,
        "MONTO_RESERVA_TEXTO": monto_a_texto_dolares(entrada["MONTO_RESERVA"]),
        "FECHA_RESERVA": entrada["FECHA_RESERVA"],
        "COMPLEMENTO_PRIMA": complemento,
        "COMPLEMENTO_PRIMA_TEXTO": monto_a_texto_dolares(entrada["COMPLEMENTO_PRIMA"]),
        "FECHA_COMPLEMENTO": entrada["FECHA_COMPLEMENTO"],
        "TERCER_PAGO": tercer_pago,
        "TERCER_PAGO_TEXTO": monto_a_texto_dolares(entrada["TERCER_PAGO"])
    }

    documento = Document(PLANTILLA)
    reemplazar_marcadores(documento, datos)

    nombre_archivo = f"Contrato_{entrada['NOMBRE'].replace(' ', '_')}.docx"
    documento.save(nombre_archivo)
    return nombre_archivo